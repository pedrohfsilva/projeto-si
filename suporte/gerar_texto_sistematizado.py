# -*- coding: utf-8 -*-
"""
Gera o Texto Sistematizado do projeto de Sistemas de Informacao (SSC0120)
- Empresa: PolymerForge 3D
- Solucao: PolyLink - Plataforma de Integracao e Rastreabilidade da Cadeia de Suprimentos

Saida: entrega/texto_sistematizado.docx
Formato: Times New Roman 12pt, entre 10 e 15 paginas de conteudo.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = REPO_ROOT / "entrega" / "texto_sistematizado.docx"


# -----------------------------------------------------------------------------
# Helpers de estilo
# -----------------------------------------------------------------------------


def _set_font(run, *, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _set_paragraph_defaults(paragraph, *, alignment=None, first_line_indent=None,
                            space_after=6, space_before=0, line_spacing=1.5):
    pf = paragraph.paragraph_format
    if alignment is not None:
        paragraph.alignment = alignment
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing


def add_paragraph(doc, text, *, size=12, bold=False, italic=False, alignment=None,
                  first_line_indent=None, space_after=6, space_before=0,
                  line_spacing=1.5, color=None):
    p = doc.add_paragraph()
    _set_paragraph_defaults(
        p,
        alignment=alignment,
        first_line_indent=first_line_indent,
        space_after=space_after,
        space_before=space_before,
        line_spacing=line_spacing,
    )
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, *, level=1):
    size = 14 if level == 1 else 12
    space_before = 14 if level == 1 else 10
    space_after = 8 if level == 1 else 6
    p = doc.add_paragraph()
    _set_paragraph_defaults(
        p,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent=None,
        space_after=space_after,
        space_before=space_before,
        line_spacing=1.15,
    )
    run = p.add_run(text)
    _set_font(run, size=size, bold=True)
    return p


def add_bullet(doc, text, *, size=12):
    p = doc.add_paragraph(style="List Bullet")
    _set_paragraph_defaults(
        p,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=None,
        space_after=4,
        space_before=0,
        line_spacing=1.3,
    )
    run = p.add_run(text)
    _set_font(run, size=size)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# -----------------------------------------------------------------------------
# Construcao do documento
# -----------------------------------------------------------------------------


def build_document() -> Document:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # =========================================================================
    # CAPA
    # =========================================================================
    add_paragraph(
        doc, "UNIVERSIDADE DE SÃO PAULO",
        size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=48, space_after=4, line_spacing=1.15,
    )
    add_paragraph(
        doc, "INSTITUTO DE CIÊNCIAS MATEMÁTICAS E DE COMPUTAÇÃO (ICMC)",
        size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4, line_spacing=1.15,
    )
    add_paragraph(
        doc, "SSC0120: Sistemas de Informação",
        size=12, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=100, line_spacing=1.15,
    )
    add_paragraph(
        doc, "PolyLink",
        size=22, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10, line_spacing=1.15,
    )
    add_paragraph(
        doc,
        "Plataforma de Integração e Rastreabilidade "
        "da Cadeia de Abastecimento para a PolymerForge 3D",
        size=14, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=80, line_spacing=1.3,
    )
    add_paragraph(
        doc, "Grupo:",
        size=12, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4, line_spacing=1.15,
    )
    integrantes = [
        "Pedro Henrique Ferreira Silva (NUSP 14677526)",
        "Enzo Tonon Morente (NUSP 14568476)",
        "Miller Matheus Lima Anacleto Rocha (NUSP 13727954)",
        "Ayrton da Costa Ganen Filho (NUSP 14560190)",
        "Matheus Rodrigues de Oliveira (NUSP 16822535)",
    ]
    for nome in integrantes:
        add_paragraph(
            doc, nome, size=12,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=2, line_spacing=1.15,
        )
    add_paragraph(
        doc, "São Carlos, julho de 2026",
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=100, line_spacing=1.15,
    )
    add_page_break(doc)

    # =========================================================================
    # SUMÁRIO
    # =========================================================================
    add_paragraph(
        doc, "SUMÁRIO",
        size=14, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18, line_spacing=1.15,
    )
    sumario = [
        ("1. Introdução", "3"),
        ("   1.1 Contextualização do problema", "3"),
        ("   1.2 Motivação e justificativa", "4"),
        ("   1.3 Soluções existentes", "4"),
        ("2. Solução Proposta", "6"),
        ("   2.1 Funcionalidades", "6"),
        ("   2.2 Questões técnicas", "7"),
        ("   2.3 Questões organizacionais", "8"),
        ("   2.4 Questões humanas", "9"),
        ("   2.5 Processos de negócio", "10"),
        ("3. Processamento da Informação", "12"),
        ("   3.1 Dados de entrada", "12"),
        ("   3.2 Processamento", "12"),
        ("   3.3 Informações geradas", "13"),
        ("4. Valor Organizacional e Vantagem Competitiva", "14"),
        ("   4.1 Cadeia de valor", "14"),
        ("   4.2 Vantagem competitiva", "14"),
        ("5. Uso de Ferramentas de Inteligência Artificial", "15"),
        ("6. Referências Bibliográficas", "17"),
    ]
    for titulo, pagina in sumario:
        p = doc.add_paragraph()
        _set_paragraph_defaults(
            p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=None, space_after=2, line_spacing=1.15,
        )
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(15), WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(f"{titulo}\t{pagina}")
        _set_font(run, size=12)
    add_page_break(doc)

    # =========================================================================
    # 1. INTRODUÇÃO
    # =========================================================================
    add_heading(doc, "1. Introdução", level=1)

    add_heading(doc, "1.1 Contextualização do problema", level=2)
    add_paragraph(
        doc,
        "A empresa-cliente deste projeto é a PolymerForge 3D, indústria "
        "brasileira sediada em São Carlos (SP) que atua desde 2019 na "
        "manufatura sob demanda de componentes plásticos para o setor "
        "automotivo. O modelo de negócio da empresa é disruptivo em dois "
        "eixos simultâneos. Primeiro, todas as peças são produzidas em "
        "impressão 3D industrial (FDM e SLS), o que permite atender "
        "encomendas de baixo volume, incluindo peças descontinuadas por "
        "montadoras, sem os custos de ferramentaria da injeção "
        "tradicional. Segundo, a operação trabalha em regime just-in-time "
        "radical: não existe estoque intermediário de polímeros; a "
        "matéria-prima é adquirida automaticamente após a confirmação de "
        "cada pedido, chega em janelas curtas de tempo e a impressão é "
        "iniciada assim que o insumo é disponibilizado.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A promessa comercial vendida aos clientes (montadoras, oficinas "
        "especializadas e retíficas) é a de entregar a peça em 48 horas "
        "após a aprovação do orçamento, sem exigência de pedido mínimo. "
        "Esse compromisso faz com que a organização dependa, de maneira "
        "crítica, do bom funcionamento das integrações com fornecedores "
        "de polímeros. Atualmente, a base tecnológica da "
        "PolymerForge é composta por três camadas: um ERP próprio, "
        "chamado IndustriaOS, que orquestra pedidos, materiais e "
        "planejamento de produção; um Portal B2B por meio do qual os "
        "clientes enviam arquivos .STL, aprovam orçamentos e acompanham "
        "seus pedidos; e um Módulo de Abastecimento Automático que, ao "
        "receber a confirmação de um pedido no ERP, dispara ordens de "
        "compra para quatro fornecedores homologados via APIs REST.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O problema selecionado pelo grupo, após análise do material de "
        "contexto da empresa e discussão do cenário, possui escopo "
        "cuidadosamente delimitado: trata-se das falhas silenciosas e da "
        "ausência de rastreabilidade nas integrações com fornecedores de "
        "matéria-prima. Não se pretende resolver, neste trabalho, as "
        "diversas outras frentes de melhoria possíveis na PolymerForge "
        "(como marketing digital, chão de fábrica ou onboarding de "
        "novos clientes), a fim de preservar o foco e o valor de "
        "negócio e evitar uma proposta ampla demais, que acabaria "
        "dispersa e sem resultado prático. O recorte escolhido é o de maior "
        "impacto financeiro imediato e aquele em que uma solução apoiada "
        "em Sistemas de Informação pode oferecer resultados mensuráveis "
        "em poucos ciclos de operação.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "1.2 Motivação e justificativa", level=2)
    add_paragraph(
        doc,
        "O problema merece atenção porque afeta diretamente o principal "
        "argumento de venda da PolymerForge: o prazo de 48 horas. As "
        "APIs dos quatro fornecedores homologados foram desenvolvidas de "
        "forma independente, sem contratos de nível de serviço "
        "formalizados e com ciclos de atualização imprevisíveis. Quando "
        "um fornecedor altera endpoints, renomeia campos ou modifica o "
        "significado de códigos de retorno, a integração do lado da "
        "PolymerForge para de funcionar corretamente. O agravante é que "
        "essas falhas não geram erros explícitos: uma requisição pode "
        "retornar HTTP 200 e, ainda assim, não ter sido processada pelo "
        "fornecedor. O ERP interpreta o código de sucesso como "
        "confirmação, atualiza o abastecimento para “material a caminho” "
        "e libera a produção para ser programada. Dois dias depois, o "
        "insumo não chega, a impressora fica ociosa e o cliente é "
        "avisado do atraso.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Os impactos são mensuráveis. Segundo o material de contexto da "
        "empresa, a taxa de entrega no prazo caiu de 94% em 2022 para "
        "71% no primeiro semestre de 2024; foram observadas, em média, "
        "2,3 paralisações não planejadas de linha por mês, com duração "
        "média de 14 horas cada; três clientes de médio porte acionaram "
        "cláusulas contratuais de penalidade por atraso; e a diretoria "
        "adiou a abertura de uma segunda unidade fabril em Campinas até "
        "que a confiabilidade operacional seja restabelecida. Em "
        "conjunto, esses indicadores tornam a resolução do problema não "
        "apenas desejável, mas condição de sobrevivência do modelo de "
        "negócio.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A solução proposta é o PolyLink, uma plataforma de integração e "
        "rastreabilidade que se coloca entre o IndustriaOS e as APIs dos "
        "fornecedores. Seus três objetivos de alto nível são: "
        "(i) desacoplar a PolymerForge das diferenças técnicas de cada "
        "fornecedor por meio de adaptadores versionados; (ii) transformar "
        "toda integração em uma transação rastreada e ativamente "
        "confirmada, encerrando o problema das falhas silenciosas; e "
        "(iii) oferecer a compras, produção e diretoria uma camada de "
        "visibilidade e decisão sobre a confiabilidade real de cada "
        "fornecedor, alimentando roteamento automático e negociações "
        "objetivas. Os benefícios esperados incluem a redução "
        "significativa das paralisações não planejadas, o retorno da "
        "taxa de entrega no prazo aos patamares de 2022 e a criação de "
        "condições mínimas para retomar o plano de expansão geográfica.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "1.3 Soluções existentes", level=2)
    add_paragraph(
        doc,
        "Há, no mercado, quatro grandes famílias de soluções que "
        "endereçam partes do problema descrito. A primeira é composta "
        "por plataformas iPaaS (Integration Platform as a Service), "
        "como MuleSoft Anypoint, Workato e Boomi, que oferecem "
        "conectores prontos e infraestrutura de integração gerenciada. "
        "A segunda família é a de ferramentas de API Gateway e "
        "observabilidade de integrações, como Kong, Apigee e Postman "
        "Monitors, focadas em roteamento, autenticação e monitoramento "
        "técnico. A terceira são os módulos avançados de Supplier "
        "Relationship Management (SRM) de ERPs de grande porte, como "
        "SAP Ariba e Oracle Procurement Cloud, orientados ao ciclo de "
        "compras completo. Por fim, há soluções verticais para cadeia "
        "de suprimentos, como Kinaxis RapidResponse e o1 TradeEdge, "
        "voltadas ao planejamento e à visibilidade multi-tier.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Cada uma dessas famílias, contudo, tem limitações quando "
        "confrontada com o cenário da PolymerForge. Plataformas iPaaS "
        "generalistas são poderosas, mas exigem investimento considerável "
        "de licenciamento, têm tempo de implementação longo e não "
        "oferecem, na configuração padrão, o conceito de “confirmação "
        "assíncrona ativa” que resolve especificamente as falhas "
        "silenciosas descritas. API Gateways são excelentes para "
        "governança técnica, porém não possuem entendimento do domínio "
        "de negócio, não classificam fornecedores por health score de "
        "abastecimento e não alimentam automaticamente decisões de "
        "roteamento comercial. Módulos SRM de ERPs corporativos são "
        "eficientes para grandes cadeias, mas subutilizam recursos em "
        "uma indústria com apenas quatro fornecedores críticos e "
        "engessam a autonomia de configuração. Soluções verticais de "
        "supply chain, por sua vez, foram desenhadas para operações com "
        "estoque de segurança, premissa que não se aplica ao just-in-time "
        "radical praticado pela empresa.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O PolyLink se diferencia por ser uma solução intermediária: "
        "leve o suficiente para ser implantada em semanas, específica o "
        "bastante para tratar o problema de falhas silenciosas em APIs "
        "de fornecedores e integrada ao domínio de negócio da manufatura "
        "sob demanda. Além disso, adiciona três elementos que nenhuma "
        "das famílias oferece nativamente: (i) confirmação assíncrona "
        "ativa baseada em callback ou polling contra o sistema-alvo do "
        "fornecedor; (ii) health score composto que combina métricas "
        "técnicas e comerciais de cada fornecedor em um único "
        "indicador; e (iii) roteamento inteligente por material, com "
        "regras editáveis por compras e produção, sem depender do time "
        "de TI a cada ajuste.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_page_break(doc)

    # =========================================================================
    # 2. SOLUÇÃO PROPOSTA
    # =========================================================================
    add_heading(doc, "2. Solução Proposta", level=1)

    add_paragraph(
        doc,
        "Antes de detalhar a solução, é útil distinguir dois tipos de "
        "sistema que costumam ser confundidos. Um ERP (Enterprise "
        "Resource Planning, ou Planejamento de Recursos Empresariais) "
        "cuida da gestão interna da empresa: registra pedidos, planeja a "
        "produção, controla materiais e reúne as informações do negócio "
        "em uma única base. Já um sistema de SCM (Supply Chain "
        "Management, ou Gestão da Cadeia de Suprimentos) cuida do fluxo "
        "de materiais e de informações entre a empresa e seus parceiros "
        "externos (fornecedores, transportadores e clientes), tratando "
        "de compras, visibilidade do abastecimento e coordenação com "
        "fornecedores. De forma resumida: o ERP olha para dentro da "
        "empresa; o SCM olha para a cadeia que a conecta ao mundo "
        "externo.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A PolymerForge já possui um ERP maduro, o IndustriaOS, mas não "
        "conta com um sistema de SCM dedicado. Hoje, as funções de "
        "cadeia de suprimentos estão embutidas de forma frágil dentro do "
        "próprio ERP, no Módulo de Abastecimento Automático, que dispara "
        "as compras mas não confirma, não registra e não acompanha o que "
        "acontece do lado do fornecedor. É justamente essa lacuna de SCM "
        "que dá origem às falhas silenciosas. O PolyLink foi concebido "
        "para preencher essa lacuna: ele não substitui o ERP (que "
        "continua sendo a fonte de verdade sobre pedidos e produção) "
        "nem pretende ser uma suíte de SCM completa (não faz, por "
        "exemplo, previsão de demanda ou gestão de estoques, funções que "
        "nem se aplicam ao modelo just-in-time radical da empresa). O "
        "PolyLink atua como uma camada especializada de execução e "
        "visibilidade da cadeia de suprimentos no elo entre a empresa e "
        "seus fornecedores, cobrindo exatamente o que o ERP não foi feito "
        "para fazer: confirmar ativamente cada compra, registrar toda a "
        "comunicação, medir a confiabilidade de cada fornecedor e rotear "
        "pedidos automaticamente.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "2.1 Funcionalidades", level=2)
    add_paragraph(
        doc,
        "O PolyLink é composto por sete funcionalidades principais, "
        "todas ilustradas no protótipo entregue com este trabalho e "
        "disponível para navegação em "
        "https://projeto-si-prototipo.vercel.app/. A "
        "escolha das funcionalidades foi guiada pelo mapeamento das duas "
        "dimensões do problema descritas no material de contexto da "
        "empresa: "
        "integrações frágeis com fornecedores e ausência de "
        "rastreabilidade e diagnóstico.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    funcionalidades = [
        (
            "Camada unificada de adaptadores",
            "cada fornecedor de polímero possui um adapter próprio, "
            "responsável por traduzir o contrato interno padronizado do "
            "PolyLink para o formato exposto pela API do fornecedor. "
            "Assim, quando um fornecedor altera endpoints, apenas seu "
            "adapter precisa ser revisado, sem impactar o ERP.",
        ),
        (
            "Confirmação assíncrona ativa",
            "para cada envio de ordem de compra, o PolyLink registra um "
            "correlation ID, aguarda o callback do fornecedor ou realiza "
            "polling periódico. Somente após a confirmação real o "
            "pedido avança no fluxo interno, eliminando a interpretação "
            "ingênua do HTTP 200.",
        ),
        (
            "Registro completo de transações (audit trail)",
            "todas as requisições, respostas, cabeçalhos, tentativas e "
            "eventos de callback são gravados com timestamp e ID de "
            "correlação, permitindo auditoria fim-a-fim de qualquer "
            "pedido em segundos.",
        ),
        (
            "Painel de health score por fornecedor",
            "o sistema calcula continuamente um score de 0 a 100 para "
            "cada fornecedor, combinando taxa de sucesso, latência p95, "
            "aderência ao contrato e cumprimento do SLA físico. Esse "
            "score alimenta o roteamento e as negociações comerciais.",
        ),
        (
            "Roteamento inteligente com fallback automático",
            "regras editáveis por compras determinam qual fornecedor "
            "deve receber cada pedido, considerando material, valor e "
            "score. Em caso de degradação ou falha, o roteamento "
            "automático redireciona o pedido para o fornecedor "
            "alternativo homologado com maior score naquele instante.",
        ),
        (
            "Central de alertas com escalação",
            "alertas classificados por severidade (P1, P2, P3) são "
            "enviados a Slack, e-mail, SMS e ao Portal B2B com políticas "
            "de escalação. Falhas silenciosas confirmadas geram alerta "
            "P1 e acionam plantão de compras em menos de um minuto.",
        ),
        (
            "Portal B2B alimentado por verdade real",
            "o Portal B2B do cliente passa a exibir o estado real do "
            "abastecimento, e não mais o estado presumido pelo ERP. Se "
            "há uma falha silenciosa em curso, o cliente é informado "
            "proativamente e recebe o novo prazo após o reroteamento.",
        ),
    ]
    for titulo, descricao in funcionalidades:
        add_bullet(doc, f"{titulo}: {descricao}")

    add_heading(doc, "2.2 Questões técnicas", level=2)
    add_paragraph(
        doc,
        "Do ponto de vista técnico, o PolyLink opera como uma camada de "
        "middleware entre o IndustriaOS e as APIs dos fornecedores, sem "
        "obrigar reescritas do ERP: o IndustriaOS continua sendo a fonte "
        "de verdade sobre pedidos e materiais e passa a delegar ao "
        "PolyLink a comunicação com fornecedores. A arquitetura "
        "recomendada é de "
        "microsserviços hospedados na mesma nuvem privada onde o "
        "IndustriaOS já opera, garantindo baixa latência entre os "
        "sistemas internos e isolamento de dados sensíveis.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Os recursos tecnológicos necessários podem ser divididos em "
        "cinco categorias. Em hardware, não há necessidade de servidores "
        "físicos adicionais: toda a solução roda sobre a mesma "
        "infraestrutura de nuvem privada já contratada. Em software, "
        "propõe-se um stack composto por um serviço de orquestração "
        "(Node.js ou Go) para os adaptadores e a máquina de estados de "
        "pedidos, um banco relacional (PostgreSQL) para dados estruturados "
        "de fornecedores, contratos e regras de roteamento, um mecanismo "
        "de fila de mensagens (RabbitMQ ou Redis Streams) para desacoplar "
        "envio e confirmação assíncrona, e um armazenamento otimizado para "
        "logs (OpenSearch ou similar) capaz de reter payloads completos "
        "por 12 meses.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Em infraestrutura, o PolyLink adota contêineres Docker orquestrados "
        "por Kubernetes, o que permite escalar horizontalmente adapters "
        "específicos em picos de demanda. Em conectividade, são "
        "estabelecidos VPNs ou túneis privados com os fornecedores que "
        "aceitarem esse padrão, aumentando a resiliência contra flutuações "
        "de Internet pública. Em observabilidade, o sistema exporta "
        "métricas para Grafana, integra-se a canais Slack, e-mail e SMS "
        "e disponibiliza um SDK simples que permite ao IndustriaOS "
        "acompanhar o estado de qualquer pedido em tempo real. Por fim, "
        "em armazenamento de dados, todas as transações são criptografadas "
        "em repouso e em trânsito, atendendo aos requisitos da LGPD e às "
        "cláusulas de sigilo dos contratos com as montadoras clientes.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Um cuidado específico foi tomado em relação à compatibilidade. "
        "Como cada fornecedor expõe uma API distinta, os adapters do "
        "PolyLink implementam um contrato interno padronizado (baseado "
        "em JSON Schema versionado) e traduzem, campo a campo, as "
        "requisições e respostas. Sempre que um fornecedor publica uma "
        "nova versão de sua API, o PolyLink é capaz de detectar a "
        "divergência de schema por meio de validação contínua contra o "
        "contrato registrado e emitir um alerta de mudança de estrutura, "
        "até mesmo antes que qualquer pedido real seja afetado. Esse "
        "mecanismo transforma um risco antes latente em uma tarefa "
        "gerenciável.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "2.3 Questões organizacionais", level=2)
    add_paragraph(
        doc,
        "A adoção do PolyLink implica mudanças em quatro frentes "
        "organizacionais. A primeira é a formalização do papel de "
        "gestor de adaptadores dentro da equipe de TI. Será necessário "
        "designar um responsável por manter as versões de contrato, "
        "revisar mudanças emitidas pelos fornecedores e coordenar as "
        "atualizações de adapter. Estima-se que essa responsabilidade "
        "possa ser exercida em regime de rotação entre os desenvolvedores "
        "da software house parceira que já atende o IndustriaOS, sem "
        "necessidade imediata de nova contratação.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A segunda frente envolve o setor de compras. Historicamente, a "
        "escolha de fornecedores era feita por relacionamento comercial "
        "e negociação informal. Com o PolyLink, esses critérios são "
        "complementados por indicadores objetivos (o health score, o "
        "histórico de rupturas e o SLA cumprido) visíveis em painel. "
        "Será preciso treinar a equipe de compras para interpretar esses "
        "indicadores e utilizá-los em conversas com fornecedores. "
        "Recomenda-se um ciclo de reuniões trimestrais com cada "
        "fornecedor, apresentando os dados do PolyLink e negociando "
        "planos de ação.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A terceira frente é a governança de regras de roteamento. Como "
        "as regras podem ser editadas por compras sem intervenção de TI, "
        "é importante estabelecer um comitê quinzenal, composto por "
        "representantes de compras, produção e TI, para revisar regras "
        "ativas, aprovar novas políticas e evitar sobreposições "
        "conflitantes. A ferramenta oferece controle de versão das "
        "regras, mas a decisão continua sendo humana. Por fim, a quarta "
        "frente diz respeito à cultura organizacional. É natural haver "
        "resistência inicial, principalmente por parte de operadores "
        "acostumados a interpretar cada alarme como falso positivo. Uma "
        "campanha interna de conscientização, associada a métricas "
        "visíveis de redução de paralisações, tende a acelerar a adoção.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Entre os riscos e barreiras à implantação, destacam-se: possível "
        "indisponibilidade de alguns fornecedores em fornecer callback "
        "assíncrono, exigindo do PolyLink o uso de polling como "
        "alternativa; complexidade em harmonizar o modelo de dados "
        "interno com contratos de APIs heterogêneas; e o custo de "
        "curadoria contínua dos adaptadores. Todos esses riscos foram "
        "considerados na arquitetura proposta, com estratégias claras "
        "de mitigação.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "2.4 Questões humanas", level=2)
    add_paragraph(
        doc,
        "A solução impacta cinco grupos de pessoas com intensidades "
        "distintas. O primeiro grupo é o de compradores e gestores de "
        "abastecimento. Para eles, o PolyLink transforma um trabalho "
        "hoje reativo (descobrir a ruptura, ligar para o fornecedor, "
        "reconstruir o histórico) em um trabalho proativo, com "
        "informações prévias, alertas precoces e apoio à decisão. "
        "Espera-se redução substancial da carga mental associada à "
        "gestão de crises e aumento da percepção de controle sobre a "
        "cadeia.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O segundo grupo são os operadores da linha de produção. "
        "Atualmente, esses profissionais aparecem no chão de fábrica "
        "para preparar uma máquina, verificam que a matéria-prima não "
        "está disponível e precisam realocar a jornada. Com o PolyLink, "
        "o agendamento das máquinas só ocorre após confirmação real do "
        "abastecimento, reduzindo o retrabalho e o desgaste. Há, ainda, "
        "ganho de segurança psicológica: o operador confia na informação "
        "exibida no painel.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O terceiro grupo é a equipe de TI e integrações. Se por um "
        "lado essa equipe assume uma nova responsabilidade (a "
        "curadoria dos adaptadores), por outro passa a contar com uma "
        "estrutura padronizada, versionada e rastreável, que reduz "
        "drasticamente o tempo médio de diagnóstico e resolução de "
        "incidentes. Um debug que hoje leva horas em ligações e "
        "planilhas passa a ser conduzido em minutos pela linha do tempo "
        "de rastreabilidade fornecida pelo sistema.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O quarto grupo é a diretoria executiva. Para os diretores, o "
        "PolyLink transforma decisões estratégicas hoje sem dados em "
        "decisões suportadas por indicadores comparáveis: qual fornecedor "
        "priorizar, qual contrato renegociar, quando expandir a base de "
        "fornecedores homologados. Isso viabiliza retomar o plano de "
        "abertura da segunda unidade fabril em Campinas com maior "
        "segurança. O quinto grupo, finalmente, são os clientes finais: "
        "montadoras, oficinas e retíficas. Eles recebem informação "
        "mais fidedigna no Portal B2B, notificações proativas em caso "
        "de reroteamento e, no médio prazo, o retorno da confiabilidade "
        "nos prazos de 48 horas que motivou o relacionamento com a "
        "PolymerForge.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "2.5 Processos de negócio", level=2)
    add_paragraph(
        doc,
        "Cinco processos de negócio da PolymerForge são afetados pela "
        "solução. O processo de compra de matéria-prima é o mais "
        "profundamente modificado. No estado atual, o Módulo de "
        "Abastecimento envia a requisição, interpreta HTTP 200 como "
        "sucesso e considera a compra concluída. No estado futuro, o "
        "processo passa a ter uma etapa formal de confirmação ativa: "
        "somente quando o fornecedor confirma o pedido de forma "
        "assíncrona (via callback ou polling), o abastecimento é "
        "marcado como confirmado e a produção é liberada. Esse ajuste, "
        "aparentemente pequeno, elimina a categoria inteira de falhas "
        "silenciosas descritas no diagnóstico.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "O processo de planejamento de produção passa a receber, do "
        "PolyLink, um sinal explícito de disponibilidade de material; "
        "hoje esse sinal é inferido erroneamente do ERP. O processo de "
        "atendimento ao cliente ganha um novo caminho de resposta "
        "proativa: em caso de reroteamento, o cliente é notificado "
        "antes que perceba a instabilidade. O processo de gestão de "
        "fornecedores, que hoje se baseia em relacionamento e "
        "negociação informal, incorpora indicadores contratuais "
        "objetivos (health score, taxa de sucesso, MTTR) e passa a "
        "ter reuniões de revisão periódicas. Por fim, o processo de "
        "resposta a incidentes é criado ou remodelado com políticas de "
        "severidade, canais de escalação e responsáveis definidos, o "
        "que traz previsibilidade onde antes havia improviso.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Os benefícios esperados desses ajustes de processo são a "
        "redução das paralisações de linha, a recuperação da taxa de "
        "entrega no prazo, a diminuição do MTTR (tempo médio de "
        "resolução de incidentes de integração) e a criação de um "
        "círculo virtuoso: quanto mais dados o PolyLink acumula, "
        "melhor fica o roteamento, menor a incidência de rupturas e "
        "maior a previsibilidade da cadeia de suprimentos.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_page_break(doc)

    # =========================================================================
    # 3. PROCESSAMENTO DA INFORMAÇÃO
    # =========================================================================
    add_heading(doc, "3. Processamento da Informação", level=1)

    add_heading(doc, "3.1 Dados de entrada", level=2)
    add_paragraph(
        doc,
        "O PolyLink consome dados provenientes de três fontes principais. "
        "A primeira é o IndustriaOS: sempre que um pedido de cliente é "
        "confirmado no Portal B2B, o ERP dispara para o PolyLink os "
        "dados do pedido: identificador, cliente, peça, especificação "
        "de material, quantidade de polímero necessária, prazo prometido "
        "e valor. Esses dados são a matriz do fluxo de abastecimento. "
        "A segunda fonte são as APIs dos fornecedores homologados: "
        "respostas síncronas às requisições, callbacks assíncronos, "
        "confirmações de despacho, notas de expedição e outros metadados "
        "sobre o pedido de compra. A terceira fonte são dados de "
        "configuração mantidos pelo próprio PolyLink: contratos de "
        "schema por fornecedor, adapters ativos, regras de roteamento, "
        "políticas de SLA e canais de alerta.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Em segundo plano, o PolyLink também consome dados operacionais "
        "auxiliares, como o calendário oficial da empresa, feriados "
        "regionais que possam afetar o transporte de insumos e "
        "informações de câmbio, relevantes para fornecedores "
        "internacionais no futuro. Todos os dados são recebidos por "
        "integrações REST ou por eventos publicados em uma fila interna "
        "de mensagens, o que garante que picos momentâneos de tráfego "
        "não gerem perda de informação.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "3.2 Processamento", level=2)
    add_paragraph(
        doc,
        "O processamento realizado pelo PolyLink combina quatro camadas "
        "lógicas. A primeira é a camada de tradução, na qual cada "
        "adapter recebe os dados internos padronizados e os traduz para "
        "o formato esperado pela API do fornecedor, e vice-versa nas "
        "respostas. A segunda é a máquina de estados de pedidos, que "
        "orquestra a transição entre os status pendente, enviado, "
        "reconhecido, em separação, em trânsito, entregue, falha e "
        "reroteado, aplicando regras de temporização e verificação a "
        "cada mudança.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A terceira camada é a analítica: com base no histórico de "
        "transações, o PolyLink calcula continuamente o health score de "
        "cada fornecedor, agrega tempos de resposta, contabiliza falhas "
        "por tipo (timeout, schema divergente, autenticação expirada, "
        "resposta ambígua) e projeta tendências por material. A quarta "
        "camada é a decisória, na qual as regras de roteamento operam "
        "sobre os indicadores em tempo real para definir qual fornecedor "
        "deve receber cada nova ordem de compra e para acionar "
        "reroteamento automático quando um fornecedor cai abaixo do "
        "threshold. Toda decisão automática é registrada com o motivo, "
        "permitindo revisão posterior por humanos.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Além disso, o PolyLink implementa regras de coerência que "
        "fazem a ponte entre a informação bruta e o significado de "
        "negócio. Uma resposta HTTP 200 sem número de rastreio, por "
        "exemplo, é sinalizada como suspeita e obriga a etapa de "
        "polling assíncrono. Uma mudança de schema detectada por diff "
        "automático contra o contrato registrado congela envios até "
        "revalidação. Uma latência p95 acima do limite acordado por "
        "30 minutos consecutivos rebaixa o fornecedor no ranking. "
        "Todas essas regras são configuráveis e auditadas.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "3.3 Informações geradas", level=2)
    add_paragraph(
        doc,
        "A partir desse processamento, o PolyLink gera quatro categorias "
        "de informação. A primeira é a informação operacional em tempo "
        "real: painel de pedidos em andamento com status real, tela de "
        "rastreabilidade fim-a-fim de cada pedido, central de alertas "
        "ativa e visão unificada da saúde das integrações. Essa camada "
        "é utilizada pela equipe de compras e produção no dia a dia.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A segunda categoria é a informação gerencial: relatórios "
        "periódicos de confiabilidade por fornecedor, indicadores "
        "consolidados de taxa de entrega no prazo, MTTR, quantidade de "
        "reroteamentos e impacto financeiro estimado de falhas "
        "evitadas. Esses relatórios apoiam as reuniões trimestrais de "
        "compras e as decisões de homologação de novos fornecedores. "
        "A terceira categoria é a informação estratégica, apresentada "
        "em dashboards sintéticos para a diretoria: evolução da "
        "confiabilidade, projeção de crescimento suportado pela cadeia "
        "e comparação entre fornecedores para negociação contratual.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "A quarta categoria é a informação voltada ao cliente. Em vez "
        "de “material a caminho” inferido a partir do ERP, o Portal B2B "
        "passa a exibir estágios reais (pedido recebido, insumo "
        "confirmado, insumo em trânsito, produção iniciada e peça a "
        "caminho), cada um deles alimentado por eventos vindos do "
        "PolyLink. Em caso de reroteamento, o cliente recebe "
        "notificação proativa e visualização da mudança de fornecedor "
        "sem impacto na especificação técnica. Assim, a informação "
        "gerada apoia a tomada de decisão em todos os níveis da "
        "organização e devolve ao cliente final a percepção de "
        "confiabilidade que era a marca de origem da PolymerForge.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_page_break(doc)

    # =========================================================================
    # 4. VALOR ORGANIZACIONAL E VANTAGEM COMPETITIVA
    # =========================================================================
    add_heading(doc, "4. Valor Organizacional e Vantagem Competitiva", level=1)

    add_heading(doc, "4.1 Cadeia de valor", level=2)
    add_paragraph(
        doc,
        "A cadeia de valor da PolymerForge 3D, adaptada do modelo "
        "clássico de Porter, é composta por atividades primárias "
        "(logística de insumos, operações de impressão 3D, logística de "
        "saída, marketing e vendas e serviço pós-venda) e por "
        "atividades de apoio como infraestrutura, gestão de pessoas, "
        "desenvolvimento tecnológico e compras. A análise do impacto do "
        "PolyLink sobre essas atividades mostra que a solução age "
        "exatamente nos pontos de atrito que hoje mais comprometem a "
        "competitividade da empresa.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Na logística de insumos, o PolyLink transforma o processo "
        "reativo em processo proativo. A confirmação assíncrona ativa "
        "elimina rupturas surpresa e o reroteamento automático mantém "
        "o fluxo mesmo diante da degradação de um fornecedor. Nas "
        "operações de impressão 3D, a garantia de que o insumo estará "
        "disponível antes de agendar a máquina reduz a ociosidade e a "
        "improvisação, aumentando o aproveitamento de horas-máquina. "
        "Na logística de saída, o Portal B2B alimentado por dados reais "
        "estabiliza expectativas de prazo e reduz o volume de "
        "atendimentos de crise. Em marketing e vendas, o retorno da "
        "promessa de 48 horas cria novamente um argumento diferenciador "
        "em propostas comerciais. No pós-venda, cai a quantidade de "
        "reclamações e o esforço de gerenciamento de contratos com "
        "cláusulas de penalidade.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Entre as atividades de apoio, compras ganha instrumentação "
        "profissional para negociar contratos, TI reduz a carga de "
        "incidentes ad hoc e passa a operar com base em contratos "
        "versionados, e gestão de pessoas tem menor pressão de "
        "turnover em setores que hoje enfrentam picos de estresse. "
        "Como consequência, a solução aumenta a eficiência, reduz "
        "custos ocultos (horas de linha ociosa, penalidades "
        "contratuais, horas de investigação manual), melhora a "
        "qualidade percebida pelo cliente e libera capital humano para "
        "iniciativas de valor superior, como novos serviços de "
        "digitalização 3D e prototipagem consultiva.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    add_heading(doc, "4.2 Vantagem competitiva", level=2)
    add_paragraph(
        doc,
        "Analisando a solução à luz das cinco forças de Porter, o "
        "PolyLink fortalece a posição da PolymerForge frente a "
        "clientes, fornecedores e concorrentes. Diante dos clientes, a "
        "solução restaura o poder de barganha da empresa ao devolver "
        "credibilidade à promessa de 48 horas, reduzindo a pressão por "
        "cláusulas de penalidade e ampliando a percepção de "
        "confiabilidade em um mercado no qual atrasos custam reputação. "
        "Diante dos fornecedores, o health score inverte a assimetria: "
        "é a PolymerForge quem passa a medir e cobrar melhorias, com "
        "dados objetivos para negociar preço, prazo e nível de serviço. "
        "Diante de novos entrantes e concorrentes diretos, o modelo "
        "just-in-time radical (hoje enfraquecido pelas rupturas) "
        "volta a ser um diferencial defensável e difícil de imitar.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_paragraph(
        doc,
        "Os benefícios estratégicos esperados podem ser sintetizados em "
        "cinco pontos. Primeiro, a recuperação da taxa de entrega no "
        "prazo, buscando retornar aos patamares anteriores à crise "
        "(94% em 2022). Segundo, a redução expressiva das "
        "paralisações não planejadas de linha. Terceiro, a viabilização "
        "financeira do plano de expansão da segunda unidade fabril em "
        "Campinas, agora sustentado por indicadores de confiabilidade. "
        "Quarto, a criação de uma capacidade defensiva que aumenta a "
        "resiliência a mudanças de fornecedores e a variações de "
        "mercado. Quinto, a formação de um ativo de dados que, ao longo "
        "do tempo, alimenta melhorias sucessivas (de modelagem "
        "preditiva de rupturas ao aprendizado sobre o comportamento de "
        "cada fornecedor), gerando vantagem competitiva sustentável.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    # =========================================================================
    # 5. USO DE IA
    # =========================================================================
    add_heading(
        doc, "5. Uso de Ferramentas de Inteligência Artificial", level=1
    )
    add_paragraph(
        doc,
        "Em atendimento à orientação expressa na descrição do projeto, "
        "declaramos com transparência as etapas em que ferramentas de "
        "inteligência artificial foram utilizadas durante o "
        "desenvolvimento deste trabalho, bem como a finalidade "
        "específica de cada uso.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )
    add_bullet(
        doc,
        "Protótipo de telas: o Claude Code (Anthropic) foi utilizado "
        "como ferramenta de apoio para a codificação do protótipo em "
        "Next.js, a partir da descrição das telas do sistema e do "
        "conjunto de funcionalidades definidos pelo grupo. A "
        "concepção das telas, a definição de fluxos, a escolha de "
        "dados exibidos e o desenho da experiência foram decisões "
        "deliberadas do grupo; a IA foi usada para acelerar a "
        "materialização dessas decisões em código.",
    )
    add_bullet(
        doc,
        "Revisão de texto: o Claude Code também foi utilizado para "
        "revisar e refinar os textos redigidos pelos integrantes do "
        "grupo, buscando maior clareza, coesão e coerência, sem "
        "alterar as ideias, argumentos e conclusões originais "
        "definidas em reuniões de grupo.",
    )
    add_paragraph(
        doc,
        "Em nenhum momento a IA foi utilizada como substituta do "
        "trabalho intelectual do grupo. A definição do problema, a "
        "escolha do recorte dentro do contexto disponível, a concepção "
        "da solução, a arquitetura proposta, o mapeamento de processos, "
        "a análise de valor organizacional e as decisões de escopo "
        "foram conduzidas pelos autores; as ferramentas serviram como "
        "apoio de produtividade em etapas de execução, sob supervisão "
        "humana contínua.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=Cm(1.25),
    )

    add_page_break(doc)

    # =========================================================================
    # 6. REFERÊNCIAS
    # =========================================================================
    add_heading(doc, "6. Referências Bibliográficas", level=1)
    referencias = [
        "LAUDON, K. C.; LAUDON, J. P. Management Information Systems: "
        "Managing the Digital Firm. 17. ed. Boston: Pearson, 2022.",
        "O'BRIEN, J. A.; MARAKAS, G. M. Administração de Sistemas de "
        "Informação. 15. ed. Porto Alegre: AMGH, 2013.",
        "PORTER, M. E. Competitive Advantage: Creating and Sustaining "
        "Superior Performance. New York: Free Press, 1985.",
        "PORTER, M. E.; MILLAR, V. E. How Information Gives You "
        "Competitive Advantage. Harvard Business Review, v. 63, n. 4, "
        "p. 149–160, 1985.",
        "CHOPRA, S.; MEINDL, P. Supply Chain Management: Strategy, "
        "Planning, and Operation. 7. ed. Boston: Pearson, 2019.",
        "TURBAN, E.; VOLONINO, L.; WOOD, G. R. Information Technology "
        "for Management: On-Demand Strategies for Performance, Growth, "
        "and Sustainability. 11. ed. Hoboken: Wiley, 2020.",
        "NEWMAN, S. Building Microservices: Designing Fine-Grained "
        "Systems. 2. ed. Sebastopol: O'Reilly, 2021.",
        "GARTNER. Market Guide for Integration Platform as a Service. "
        "Stamford: Gartner Inc., 2024.",
        "MULE SOFT. State of API and Integration Report. San Francisco: "
        "MuleSoft, 2024.",
        "ANTHROPIC. Claude Code Documentation. 2026. Disponível em: "
        "https://docs.anthropic.com/claude-code. Acesso em: jul. 2026.",
        "POLYMERFORGE 3D. Documento de Contexto e Problema. Material "
        "interno da disciplina SSC0120, ICMC-USP, 2026.",
    ]
    for ref in referencias:
        p = doc.add_paragraph()
        _set_paragraph_defaults(
            p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=None, space_after=6, line_spacing=1.15,
        )
        run = p.add_run(ref)
        _set_font(run, size=12)

    return doc


def main():
    doc = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Documento gerado em: {OUTPUT}")


if __name__ == "__main__":
    main()
